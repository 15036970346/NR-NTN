#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""把 NTN 全链路整合计划 (markdown) 渲染成结构化 Word 文档。"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

CN = '宋体'
MONO = 'Consolas'

doc = Document()

# 默认正文字体 (含中文 eastAsia)
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), CN)


def style_runs(par, name=CN):
    for r in par.runs:
        rpr = r._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            from docx.oxml import OxmlElement
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:eastAsia'), name)


def add_heading(text, level):
    h = doc.add_heading(text, level)
    style_runs(h)
    return h


def add_code(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    r = p.add_run('\n'.join(lines))
    r.font.name = MONO
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), MONO)
    # 浅灰底
    shd = p._p.get_or_add_pPr()
    from docx.oxml import OxmlElement
    e = OxmlElement('w:shd')
    e.set(qn('w:val'), 'clear')
    e.set(qn('w:fill'), 'F2F2F2')
    shd.append(e)


# 粗体行内 **...**
def add_para(text):
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for seg in parts:
        if not seg:
            continue
        if seg.startswith('**') and seg.endswith('**'):
            r = p.add_run(seg[2:-2]); r.bold = True
        else:
            r = p.add_run(seg)
    style_runs(p)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for seg in parts:
        if not seg:
            continue
        if seg.startswith('**') and seg.endswith('**'):
            r = p.add_run(seg[2:-2]); r.bold = True
        else:
            r = p.add_run(seg)
    style_runs(p)


def add_table(rows):
    cols = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(cell)
            if i == 0:
                r.bold = True
            r.font.size = Pt(9)
            style_runs(p)


SRC = '/home/cgf/.claude/plans/contrib-geo-sat-examples-ntn-system-sim-quizzical-beacon.md'
with open(SRC, encoding='utf-8') as f:
    md = f.read().splitlines()

i = 0
n = len(md)
while i < n:
    line = md[i]
    # 代码块
    if line.strip().startswith('```'):
        buf = []
        i += 1
        while i < n and not md[i].strip().startswith('```'):
            buf.append(md[i]); i += 1
        add_code(buf)
        i += 1
        continue
    # 表格
    if line.strip().startswith('|') and i + 1 < n and re.match(r'^\s*\|[-:\s|]+\|\s*$', md[i+1]):
        rows = []
        def cells(l):
            return [c.strip() for c in l.strip().strip('|').split('|')]
        rows.append(cells(line))
        i += 2  # skip separator
        while i < n and md[i].strip().startswith('|'):
            rows.append(cells(md[i])); i += 1
        add_table(rows)
        continue
    m = re.match(r'^(#{1,6})\s+(.*)$', line)
    if m:
        add_heading(m.group(2), min(len(m.group(1)), 4))
        i += 1
        continue
    if re.match(r'^\s*[-*]\s+', line):
        add_bullet(re.sub(r'^\s*[-*]\s+', '', line))
        i += 1
        continue
    if re.match(r'^\s*\d+\.\s+', line):
        add_para(line.strip())
        i += 1
        continue
    if line.strip() == '---':
        i += 1
        continue
    if line.strip() == '':
        i += 1
        continue
    add_para(line)
    i += 1

OUT = '/home/cgf/Downloads/ns-3-dev/contrib/geo-sat/NTN全链路整合实现方案.docx'
doc.save(OUT)
print('saved:', OUT)
